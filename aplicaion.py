import streamlit as st
import pdfplumber
import pandas as pd
import re
import unicodedata
import copy
import subprocess
import tempfile
import os
import zipfile
import pypdf
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

def verificar_password():
    """Muestra un formulario de contraseña y detiene la app si es incorrecta."""
    def password_entered():
        if st.session_state["password"] == st.secrets["password"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.text_input(
            "🔒 Contraseña de acceso", type="password",
            on_change=password_entered, key="password"
        )
        st.stop()
    elif not st.session_state["password_correct"]:
        st.text_input(
            "🔒 Contraseña de acceso", type="password",
            on_change=password_entered, key="password"
        )
        st.error("❌ Contraseña incorrecta")
        st.stop()

verificar_password()



_LOGO_SVG = """<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100" width="90" height="90">
  <rect width="100" height="100" rx="14" fill="#1a1f2e" stroke="rgba(57,211,83,0.4)" stroke-width="2"/>
  <circle cx="50" cy="38" r="22" fill="none" stroke="#39d353" stroke-width="3.5"/>
  <circle cx="50" cy="38" r="14" fill="#39d353"/>
  <rect x="28" y="64" width="44" height="6" rx="3" fill="#39d353"/>
  <rect x="34" y="74" width="32" height="5" rx="2.5" fill="#39d353" opacity="0.7"/>
  <text x="50" y="96" text-anchor="middle" font-family="Barlow,Arial,sans-serif"
        font-weight="900" font-size="9" fill="#39d353" letter-spacing="2">SENA</text>
</svg>"""

st.set_page_config(page_title="Automatizador SENA", page_icon="🎓", layout="wide")

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Barlow:wght@400;600;700;900&family=Barlow+Condensed:wght@700;900&display=swap');

html, body, [data-testid="stAppViewContainer"] {{
    background-color: #0d1117 !important;
    color: #e6edf3 !important;
    font-family: 'Barlow', sans-serif !important;
}}
[data-testid="stHeader"]  {{ background-color: #0d1117 !important; }}
[data-testid="stSidebar"] {{ background-color: #161b22 !important; }}

.sena-header {{
    display:flex; flex-direction:column; align-items:center; justify-content:center;
    padding:2.2rem 1rem 1.4rem; margin-bottom:0.5rem;
    background:linear-gradient(135deg,#0d1117 0%,#161b22 60%,#0d2137 100%);
    border-bottom:3px solid #39d353; position:relative; overflow:hidden;
}}
.sena-header::before {{
    content:''; position:absolute; top:-60px; left:-60px;
    width:200px; height:200px; border-radius:50%;
    background:radial-gradient(circle,rgba(57,211,83,.12) 0%,transparent 70%);
    pointer-events:none;
}}
.sena-header::after {{
    content:''; position:absolute; bottom:-60px; right:-60px;
    width:200px; height:200px; border-radius:50%;
    background:radial-gradient(circle,rgba(57,211,83,.08) 0%,transparent 70%);
    pointer-events:none;
}}
.sena-logo-wrap  {{ display:flex; align-items:center; gap:1.4rem; z-index:1; }}
.sena-logo-img   {{
    height:90px; width:90px;
}}
.sena-logo-badge {{ display:flex; flex-direction:column; align-items:flex-start; gap:.2rem; }}
.sena-logo-sub   {{
    font-family:'Barlow',sans-serif; font-weight:700; font-size:.8rem;
    letter-spacing:.18em; color:#8b949e; text-transform:uppercase;
}}
.sena-title {{
    font-family:'Barlow Condensed',sans-serif; font-weight:900; font-size:1.7rem;
    letter-spacing:.06em; color:#e6edf3; text-transform:uppercase;
    margin-top:.9rem; text-align:center; z-index:1;
}}
.sena-subtitle {{ font-size:.9rem; color:#8b949e; margin-top:.3rem; text-align:center; z-index:1; }}
.divider {{
    height:1px; background:linear-gradient(90deg,transparent,#30363d,transparent);
    margin:1.2rem 0;
}}
[data-testid="stFileUploader"] {{
    background:#161b22 !important; border:1.5px dashed #30363d !important;
    border-radius:10px !important; padding:.5rem !important; transition:border-color .2s;
}}
[data-testid="stFileUploader"]:hover {{ border-color:#39d353 !important; }}
[data-testid="stFileUploader"] label {{
    color:#e6edf3 !important; font-family:'Barlow',sans-serif !important; font-weight:600 !important;
}}
.stButton>button {{
    background:linear-gradient(135deg,#238636,#2ea043) !important;
    color:#fff !important; font-family:'Barlow',sans-serif !important;
    font-weight:700 !important; font-size:1rem !important; letter-spacing:.04em !important;
    border:none !important; border-radius:8px !important; padding:.65rem 1.8rem !important;
    box-shadow:0 4px 16px rgba(46,160,67,.35) !important; transition:all .2s !important;
    text-transform:uppercase !important;
}}
.stButton>button:hover {{
    background:linear-gradient(135deg,#2ea043,#3fb950) !important;
    box-shadow:0 6px 24px rgba(63,185,80,.45) !important; transform:translateY(-1px) !important;
}}
[data-testid="stDownloadButton"]>button {{
    background:linear-gradient(135deg,#1f6feb,#388bfd) !important;
    box-shadow:0 4px 16px rgba(56,139,253,.35) !important;
}}
[data-testid="stDownloadButton"]>button:hover {{
    background:linear-gradient(135deg,#388bfd,#58a6ff) !important;
}}
[data-testid="stSuccess"]  {{ background:rgba(46,160,67,.13)  !important; border-left:4px solid #2ea043 !important; color:#56d364 !important; border-radius:8px !important; }}
[data-testid="stError"]    {{ background:rgba(248,81,73,.13)  !important; border-left:4px solid #f85149 !important; color:#ff7b72 !important; border-radius:8px !important; }}
[data-testid="stWarning"]  {{ background:rgba(210,153,34,.13) !important; border-left:4px solid #d29922 !important; color:#e3b341 !important; border-radius:8px !important; }}
[data-testid="stInfo"]     {{ background:rgba(56,139,253,.13) !important; border-left:4px solid #388bfd !important; color:#79c0ff !important; border-radius:8px !important; }}
[data-testid="stDataFrame"] {{
    background:#161b22 !important; border:1px solid #30363d !important;
    border-radius:10px !important; overflow:hidden !important;
}}
.stProgress>div>div>div>div {{
    background:linear-gradient(90deg,#238636,#39d353) !important; border-radius:4px !important;
}}
h1,h2,h3 {{
    color:#e6edf3 !important; font-family:'Barlow Condensed',sans-serif !important;
    font-weight:700 !important; letter-spacing:.03em !important;
}}
p,li,label {{ color:#c9d1d9 !important; font-family:'Barlow',sans-serif !important; }}
.seccion-card {{
    background:#161b22; border:1px solid #30363d;
    border-radius:12px; padding:1.4rem 1.6rem; margin-bottom:1rem;
}}
.seccion-card-title {{
    font-family:'Barlow Condensed',sans-serif; font-weight:700; font-size:1.05rem;
    color:#39d353; letter-spacing:.08em; text-transform:uppercase; margin-bottom:.6rem;
}}
.chip-row {{ display:flex; gap:.8rem; flex-wrap:wrap; margin:.8rem 0; }}
.chip {{
    padding:.35rem 1rem; border-radius:20px; font-family:'Barlow',sans-serif;
    font-weight:700; font-size:.85rem; letter-spacing:.04em;
}}
.chip-green  {{ background:rgba(46,160,67,.2);  color:#56d364; border:1px solid #238636; }}
.chip-red    {{ background:rgba(248,81,73,.2);  color:#ff7b72; border:1px solid #f85149; }}
.chip-blue   {{ background:rgba(56,139,253,.2); color:#79c0ff; border:1px solid #388bfd; }}
.chip-yellow {{ background:rgba(210,153,34,.2); color:#e3b341; border:1px solid #d29922; }}
.chip-purple {{ background:rgba(163,113,247,.2); color:#d2a8ff; border:1px solid #8957e5; }}
.match-box {{
    background:#1c2128; border:1px solid #30363d; border-radius:8px;
    padding:.7rem 1rem; margin:.3rem 0; font-size:.88rem; color:#c9d1d9;
}}
.match-high {{ border-left:4px solid #39d353; }}
.match-med  {{ border-left:4px solid #e3b341; }}
.match-low  {{ border-left:4px solid #ff7b72; }}
.split-row {{
    display:flex; align-items:center; gap:1rem;
    background:#1c2128; border:1px solid #30363d;
    border-radius:8px; padding:.6rem 1rem; margin:.25rem 0;
    font-size:.88rem; color:#c9d1d9;
}}
.split-row-ok   {{ border-left:4px solid #39d353; }}
.split-row-warn {{ border-left:4px solid #e3b341; }}
.split-row-err  {{ border-left:4px solid #f85149; }}
.split-num {{
    font-family:'Barlow Condensed',sans-serif; font-weight:900;
    font-size:1rem; color:#8b949e; min-width:2rem; text-align:right;
}}
</style>

<div class="sena-header">
  <div class="sena-logo-wrap">
    {_LOGO_SVG}
    <div class="sena-logo-badge">
      <span class="sena-logo-sub">Servicio Nacional de Aprendizaje</span>
       <div class="sena-title">  FRIJOLITO</div>
    </div>
  </div>
  <div class="sena-title">Automatizador de Obligaciones CIAA</div>
  <div class="sena-subtitle">Sistema automático para copiar tablas de obligaciones desde PDFs hacia el Word combinado</div>
</div>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
#  UTILIDADES COMPARTIDAS
# ══════════════════════════════════════════════════════════════════

_STOPWORDS_NOMBRES = {
    'DE', 'DEL', 'LA', 'EL', 'LOS', 'LAS', 'Y', 'E', 'A',
    'SAN', 'SANTA', 'VON', 'VAN', 'DOS', 'DA', 'DI', 'MC', 'MAC'
}


def normalizar(texto: str) -> str:
    if not texto:
        return ""
    texto = unicodedata.normalize('NFD', str(texto))
    texto = ''.join(c for c in texto if unicodedata.category(c) != 'Mn')
    texto = re.sub(r'\s+', ' ', texto)
    return texto.upper().strip()


def _tokens(nombre_norm: str) -> list:
    return [t for t in nombre_norm.split()
            if len(t) >= 3 and t not in _STOPWORDS_NOMBRES]


def score_nombre(a: str, b: str) -> int:
    if not a or not b:
        return 0
    na, nb = normalizar(a), normalizar(b)
    if na == nb:
        return 100
    if na in nb or nb in na:
        return 98

    ta, tb = _tokens(na), _tokens(nb)
    if not ta or not tb:
        return 0

    sa, sb = set(ta), set(tb)
    interseccion_exacta = sa & sb
    menor_e = sa if len(sa) <= len(sb) else sb
    mayor_e = sb if len(sa) <= len(sb) else sa

    if menor_e and menor_e.issubset(mayor_e):
        return 96

    jaccard = len(interseccion_exacta) / len(sa | sb)
    if jaccard >= 0.60:
        return int(85 + jaccard * 10)

    cobertura_exacta = len(interseccion_exacta) / len(menor_e) if menor_e else 0
    if len(interseccion_exacta) >= 2 and cobertura_exacta >= 0.60:
        return 75

    def distancia_levenshtein(s1, s2):
        if s1 == s2:
            return 0
        if len(s1) < len(s2):
            s1, s2 = s2, s1
        if not s2:
            return len(s1)
        prev = list(range(len(s2) + 1))
        for i, c1 in enumerate(s1):
            curr = [i + 1]
            for j, c2 in enumerate(s2):
                curr.append(min(prev[j + 1] + 1, curr[j] + 1,
                                prev[j] + (0 if c1 == c2 else 1)))
            prev = curr
        return prev[-1]

    def tokens_similares(set_a, set_b, max_dist=1):
        pares = set()
        for ta_tok in set_a:
            for tb_tok in set_b:
                if distancia_levenshtein(ta_tok, tb_tok) <= max_dist:
                    pares.add(ta_tok)
                    break
        return pares

    sim_fuzzy = tokens_similares(sa, sb)
    menor_f = sa if len(sa) <= len(sb) else sb
    mayor_f = sb if len(sa) <= len(sb) else sa

    sim_en_mayor = tokens_similares(menor_f, mayor_f)
    if (sim_en_mayor and sim_en_mayor.issuperset(menor_f - sim_en_mayor) or
            len(sim_en_mayor) == len(menor_f)):
        return 92

    cobertura_fuzzy = len(sim_fuzzy) / len(menor_f) if menor_f else 0
    if len(sim_fuzzy) >= 2 and cobertura_fuzzy >= 0.60:
        return 60

    return 0


UMBRAL = 70


def buscar_mejor_match(nombre_pdf: str, mapa_word: dict):
    mejor_nombre = None
    mejor_score = 0
    for nw in mapa_word:
        s = score_nombre(nombre_pdf, nw)
        if s > mejor_score:
            mejor_score = s
            mejor_nombre = nw
    if mejor_score >= UMBRAL:
        return mejor_nombre, mejor_score
    return None, mejor_score


def texto_parrafo(elem) -> str:
    return ''.join(r.text for r in elem.iter()
                   if r.tag.endswith('}t') and r.text).strip()


# ══════════════════════════════════════════════════════════════════
#  MÓDULO 1 — COPIAR TABLAS DE OBLIGACIONES (lógica original)
# ══════════════════════════════════════════════════════════════════

def extraer_nombre_pdf(archivo_pdf):
    try:
        archivo_pdf.seek(0)
        with pdfplumber.open(archivo_pdf) as pdf:
            texto_completo = ""
            for p in pdf.pages[:2]:
                texto_completo += (p.extract_text() or "") + " "

        tn = normalizar(texto_completo)

        m = re.search(
            r'\bCONTRATISTA\s+([A-ZÁÉÍÓÚÑ ]{8,80}?)\s+(?:FECHA|OBJETO|VALOR|INICIO)',
            tn)
        if m:
            return re.sub(r'\s+', ' ', m.group(1)).strip()

        m = re.search(r'([A-ZÁÉÍÓÚÑ ]{8,80})\s*,?\s*IDENTIFICAD[OA]\s+CON', tn)
        if m:
            nombre = re.sub(r'\s+', ' ', m.group(1)).strip()
            return re.sub(r'^(EL|LA|DE|DEL|Y|EN)\s+', '', nombre).strip()

        m = re.search(
            r'([A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ]+(?: [A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ]+){1,5})'
            r',?\s+identificad[oa]\s+con',
            texto_completo, re.IGNORECASE)
        if m:
            return normalizar(m.group(1))

        return None
    except Exception:
        return None


def extraer_tabla_obligaciones_pdf(archivo_pdf):
    STOP_KW = [
        '¿EL CONTRATISTA', 'CONCLUSIONES', 'RECOMENDACIONES',
        'VERIFICACION', 'DESPLAZAMIENTO', 'ITEM NRO'
    ]

    def es_stop(celda):
        return any(normalizar(k) in normalizar(str(celda or "")) for k in STOP_KW)

    filas_crudas = []
    header = None
    header_encontrado = False

    try:
        archivo_pdf.seek(0)
        with pdfplumber.open(archivo_pdf) as pdf:
            for pagina in pdf.pages:
                for tabla in pagina.extract_tables():
                    if not tabla:
                        continue
                    primera = str(tabla[0][0] or "").strip()
                    es_encabezado = (
                        re.match(r'^N[rRoO°º\.]+\s*$', primera) is not None
                        or primera in ("Nro.", "No.", "Nro", "No", "N°", "Nº")
                        or (len(tabla[0]) >= 2
                            and 'BLIGACION' in normalizar(str(tabla[0][1] or '')))
                    )
                    if es_encabezado and not header_encontrado:
                        header = tabla[0]
                        header_encontrado = True
                        for fila in tabla[1:]:
                            if es_stop(fila[0]):
                                break
                            filas_crudas.append(fila)
                        continue
                    if header_encontrado:
                        if es_stop(tabla[0][0]):
                            continue
                        cel0 = str(tabla[0][0] or "").strip()
                        if cel0 == '' or cel0.isdigit():
                            for fila in tabla:
                                if es_stop(fila[0]):
                                    break
                                filas_crudas.append(fila)

        if not header or not filas_crudas:
            return None

        obligaciones = []
        obligacion_actual = None

        for fila in filas_crudas:
            nro = str(fila[0] or "").strip()
            cols = [str(fila[i] or "").strip() if i < len(fila) else "" for i in range(1, 5)]
            if nro.isdigit():
                if obligacion_actual:
                    obligaciones.append(obligacion_actual)
                obligacion_actual = [nro] + cols
            else:
                if obligacion_actual:
                    for i, c in enumerate(cols, start=1):
                        if c:
                            obligacion_actual[i] = (obligacion_actual[i] + " " + c).strip()

        if obligacion_actual:
            obligaciones.append(obligacion_actual)

        return [header] + obligaciones

    except Exception:
        return None


def construir_mapa_word(doc) -> dict:
    NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body
    children = list(body)

    nombre_en_pos = {}
    for i, elem in enumerate(children):
        if elem.tag.split('}')[-1] != 'tbl':
            continue
        try:
            for fila in elem.findall(f'.//{{{NS}}}tr'):
                celdas = fila.findall(f'.//{{{NS}}}tc')
                if len(celdas) < 2:
                    continue
                c0 = normalizar(''.join(
                    t.text for t in celdas[0].iter()
                    if t.tag.endswith('}t') and t.text))
                if 'CONTRATISTA' in c0:
                    nombre = normalizar(''.join(
                        t.text for t in celdas[1].iter()
                        if t.tag.endswith('}t') and t.text))
                    if nombre:
                        nombre_en_pos[i] = nombre
                    break
        except Exception:
            continue

    INICIO_KW = [
        'OBLIGACIONES ESPECIFICAS DEL CONTRATISTA',
        'OBLIGACIONES ESPECIFICAS',
        'OBLIGACIONES DEL CONTRATISTA',
    ]
    FIN_KW = [
        'VERIFICACION DE ANTECEDENTES',
        'VERIFICACION DE ANTECEDENTE',
        'VERIFICACION ANTECEDENTES',
    ]

    mapa = {}
    dentro = False
    tablas_bloque = []
    pos_inicio = None
    nombre_bloque = None

    for i, elem in enumerate(children):
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            txt = normalizar(texto_parrafo(elem))
            if not dentro and any(k in txt for k in INICIO_KW):
                dentro = True
                pos_inicio = i
                tablas_bloque = []
                nombre_bloque = None
                for pos, nom in sorted(nombre_en_pos.items(), reverse=True):
                    if pos < i:
                        nombre_bloque = nom
                        break
            elif dentro and any(k in txt for k in FIN_KW):
                if nombre_bloque and tablas_bloque:
                    mapa[nombre_bloque] = {
                        "tablas": tablas_bloque.copy(),
                        "inicio": pos_inicio,
                        "fin": i
                    }
                dentro = False
                tablas_bloque = []
        elif tag == 'tbl' and dentro:
            tablas_bloque.append(elem)

    if dentro and nombre_bloque and tablas_bloque:
        mapa[nombre_bloque] = {
            "tablas": tablas_bloque.copy(),
            "inicio": pos_inicio,
            "fin": len(children)
        }

    return mapa


def reemplazar_tablas_obligaciones(doc, info_bloque: dict, datos_pdf: list) -> int:
    tablas_xml = info_bloque["tablas"]
    if not tablas_xml:
        return 0

    tabla_destino = None
    for t in doc.tables:
        if t._element is tablas_xml[0]:
            tabla_destino = t
            break

    if tabla_destino is None:
        raise Exception("No se encontró la tabla destino en el documento.")

    idx_enc = 0
    ncols_word = len(tabla_destino.columns)

    for i, fila in enumerate(tabla_destino.rows):
        txt = normalizar(" ".join(c.text for c in fila.cells))
        if any(k in txt for k in ('OBLIGACION', 'ACCIONES', 'EVIDENCIAS', 'ACTIVIDAD')):
            idx_enc = i
            break

    for fi in reversed(range(idx_enc + 1, len(tabla_destino.rows))):
        tabla_destino._element.remove(tabla_destino.rows[fi]._tr)

    for fila_pdf in datos_pdf[1:]:
        nuevas_celdas = tabla_destino.add_row().cells
        for j in range(ncols_word):
            valor = str(fila_pdf[j] or "").strip() if j < len(fila_pdf) else ""
            nuevas_celdas[j].text = valor

    for tabla_xml in tablas_xml[1:]:
        try:
            tabla_xml.getparent().remove(tabla_xml)
        except Exception:
            pass

    return len(tablas_xml)


# ══════════════════════════════════════════════════════════════════
#  MÓDULO 2 — DIVISOR DE INFORMES (existente)
# ══════════════════════════════════════════════════════════════════

def encontrar_bloques_informes(doc) -> list:
    """
    Detecta cada informe individual dentro del combinado.
    Inicio real: tabla de encabezado (PROCESO/GESTIÓN CONTRACTUAL) que precede al marcador.
    Marcador de referencia: párrafo con 'INFORME DE SUPERVISIÓN' + 'CONTRATO NRO'
    Nombre: primera celda CONTRATISTA dentro del bloque.
    Retorna lista de dicts: {nombre, inicio, fin}
    """
    NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body
    children = list(body)

    _FALSOS_CONTRATISTA = {'ABRIL', 'MAYO', 'JUNIO', 'JULIO', 'AGOSTO',
                           'SEPTIEMBRE', 'OCTUBRE', 'NOVIEMBRE', 'DICIEMBRE',
                           'ENERO', 'FEBRERO', 'MARZO'}

    KW_ENCABEZADO = {'PROCESO', 'GESTION CONTRACTUAL', 'NOMBRE DEL FORMATO',
                     'INFORME DE SUPERVISION', 'CLASIFICACION DE LA INFORMACION'}

    def es_tabla_encabezado(elem):
        """True si la tabla contiene las palabras clave del encabezado SENA."""
        txt = normalizar(texto_parrafo(elem) or
                         ' '.join(t.text or '' for t in elem.iter() if t.tag.endswith('}t')))
        hits = sum(1 for kw in KW_ENCABEZADO if kw in txt)
        return hits >= 2

    bloques = []

    for i, elem in enumerate(children):
        if elem.tag.split('}')[-1] != 'p':
            continue
        txt = normalizar(texto_parrafo(elem))
        if 'INFORME DE SUPERVISION' in txt and 'CONTRATO NRO' in txt:
            # Buscar hacia atrás la tabla de encabezado más cercana
            inicio_real = i  # fallback: desde el marcador mismo
            for j in range(i - 1, max(i - 6, -1), -1):
                tag_j = children[j].tag.split('}')[-1]
                if tag_j == 'tbl' and es_tabla_encabezado(children[j]):
                    # Incluir párrafos vacíos antes de la tabla
                    inicio_real = j
                    # Retroceder más para incluir párrafos vacíos previos
                    while inicio_real > 0 and children[inicio_real - 1].tag.split('}')[-1] == 'p'                             and not texto_parrafo(children[inicio_real - 1]).strip():
                        inicio_real -= 1
                    break

            # Cerrar bloque anterior justo antes del nuevo inicio
            if bloques:
                bloques[-1]['fin'] = inicio_real

            bloques.append({'nombre': None, 'inicio': inicio_real, 'fin': len(children)})

    # Buscar nombre CONTRATISTA en cada bloque
    for bloque in bloques:
        ini = bloque['inicio']
        fin = bloque['fin']
        for elem in children[ini:fin]:
            if elem.tag.split('}')[-1] != 'tbl':
                continue
            for fila in elem.findall(f'.//{{{NS}}}tr'):
                celdas = fila.findall(f'.//{{{NS}}}tc')
                if len(celdas) < 2:
                    continue
                c0 = normalizar(''.join(
                    t.text for t in celdas[0].iter() if t.tag.endswith('}t') and t.text))
                if 'CONTRATISTA' in c0:
                    nombre_raw = ''.join(
                        t.text for t in celdas[1].iter()
                        if t.tag.endswith('}t') and t.text).strip()
                    nombre_up = nombre_raw.upper().strip()
                    if (nombre_raw and len(nombre_raw) > 5
                            and nombre_up not in _FALSOS_CONTRATISTA
                            and not nombre_up.startswith('SENA')
                            and not nombre_up.startswith('SERVICIO')):
                        bloque['nombre'] = nombre_raw
                        break
            if bloque['nombre']:
                break

    return bloques


def slug_nombre(nombre: str) -> str:
    """Convierte un nombre a formato seguro para nombre de archivo."""
    n = normalizar(nombre)
    n = re.sub(r'[^A-Z0-9 ]', '', n)
    n = re.sub(r'\s+', '_', n.strip())
    return n


def _extraer_xml_combinado(doc_origen) -> bytes:
    """
    Extrae el XML del body del documento combinado directamente
    desde los bytes del archivo original guardado en session_state.
    Se usa junto con los bytes originales para clonar el ZIP.
    """
    return doc_origen.element.body


def clonar_bloque_como_zip(original_zip_bytes: bytes, inicio: int, fin: int) -> bytes:
    """
    Clona el ZIP completo del docx original pero reemplazando el body
    del document.xml para contener solo los elementos [inicio:fin].
    Esto preserva 100% imágenes, estilos, headers, temas y fuentes.
    """
    NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    with zipfile.ZipFile(BytesIO(original_zip_bytes)) as z:
        orig_doc_xml = z.read('word/document.xml')

    # Parsear y manipular el body
    tree = etree.fromstring(orig_doc_xml)
    body = tree.find(f'{{{NS}}}body')
    children = list(body)

    # Guardar sectPr (configuración de página)
    sect_pr = None
    if children and children[-1].tag.split('}')[-1] == 'sectPr':
        sect_pr = children[-1]

    # Vaciar body
    for elem in list(body):
        body.remove(elem)

    # Insertar solo el rango del informe
    for elem in children[inicio:fin]:
        body.append(elem)

    # Restaurar sectPr
    if sect_pr is not None:
        body.append(sect_pr)

    new_doc_xml = etree.tostring(
        tree, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Clonar ZIP completo reemplazando solo document.xml
    out_buf = BytesIO()
    with zipfile.ZipFile(BytesIO(original_zip_bytes)) as zin:
        with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == 'word/document.xml':
                    zout.writestr(item, new_doc_xml)
                else:
                    zout.writestr(item, zin.read(item.filename))

    out_buf.seek(0)
    return out_buf.getvalue()


def _buscar_libreoffice() -> str | None:
    """
    Detecta el ejecutable de LibreOffice según el sistema operativo.
    Retorna la ruta al ejecutable o None si no está instalado.
    """
    import platform
    sistema = platform.system()

    if sistema == "Windows":
        rutas_posibles = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for ruta in rutas_posibles:
            if os.path.exists(ruta):
                return ruta
        # Intentar desde PATH
        import shutil
        return shutil.which("soffice") or shutil.which("libreoffice")

    else:  # Linux / Mac
        import shutil
        return shutil.which("libreoffice") or shutil.which("soffice")


def docx_a_pdf_libreoffice(docx_bytes: bytes, nombre_base: str) -> bytes | None:
    """
    Convierte bytes de un .docx a bytes de .pdf usando LibreOffice headless.
    Detecta automáticamente la ruta en Windows y Linux/Mac.
    Retorna bytes del PDF o None si falla.
    """
    ejecutable = _buscar_libreoffice()
    if not ejecutable:
        return None  # LibreOffice no instalado

    with tempfile.TemporaryDirectory() as tmpdir:
        ruta_docx = os.path.join(tmpdir, f"{nombre_base}.docx")
        ruta_pdf  = os.path.join(tmpdir, f"{nombre_base}.pdf")

        with open(ruta_docx, 'wb') as f:
            f.write(docx_bytes)

        result = subprocess.run(
            [ejecutable, '--headless', '--convert-to', 'pdf',
             '--outdir', tmpdir, ruta_docx],
            capture_output=True, text=True, timeout=120
        )

        if result.returncode == 0 and os.path.exists(ruta_pdf):
            with open(ruta_pdf, 'rb') as f:
                return f.read()
        return None


def generar_zip_informes(original_zip_bytes: bytes, bloques: list, seleccionados: list,
                         formato: str, progreso_widget) -> bytes:
    """
    Genera un ZIP con los informes individuales en el formato indicado ('docx' o 'pdf').
    Usa clonar_bloque_como_zip para preservar 100% el formato original
    (imágenes, estilos, encabezados, temas, fuentes).
    Solo incluye los bloques cuyo nombre está en `seleccionados`.
    """
    zip_buf = BytesIO()
    total = len(seleccionados)

    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        contador = 0
        for idx, bloque in enumerate(bloques):
            nombre = bloque.get('nombre') or f"SIN_NOMBRE_{idx+1}"
            if nombre not in seleccionados:
                continue

            contador += 1
            slug = slug_nombre(nombre)
            nombre_archivo = f"{contador:02d}_{slug}"

            progreso_widget.progress(
                contador / total,
                text=f"Generando {contador}/{total}: {nombre}"
            )

            # Clonar ZIP completo con solo este bloque — preserva todo el formato
            docx_bytes = clonar_bloque_como_zip(
                original_zip_bytes, bloque['inicio'], bloque['fin'])

            if formato == 'docx':
                zf.writestr(f"{nombre_archivo}.docx", docx_bytes)
            else:  # pdf
                pdf_bytes = docx_a_pdf_libreoffice(docx_bytes, nombre_archivo)
                if pdf_bytes:
                    zf.writestr(f"{nombre_archivo}.pdf", pdf_bytes)
                else:
                    zf.writestr(f"{nombre_archivo}_CONVERSION_FALLIDA.docx", docx_bytes)

    zip_buf.seek(0)
    return zip_buf.getvalue()


# ══════════════════════════════════════════════════════════════════
#  MÓDULO 3 — UNIR PDFS POR CÉDULA (nuevo)
# ══════════════════════════════════════════════════════════════════

def extraer_cedula_de_nombre_archivo(nombre_archivo: str):
    """
    Busca una secuencia de 6 a 12 dígitos en el nombre del archivo.
    Como estos certificados siempre se descargan y guardan usando la
    cédula (ej. '42784175_PROCURADURIA.pdf'), esta es la forma MÁS
    confiable de identificar a quién pertenece cada PDF.
    """
    m = re.search(r'(\d{6,12})', nombre_archivo)
    return m.group(1) if m else None


def extraer_persona_pdf(archivo_pdf):
    """
    Extrae cédula y nombre analizando el CONTENIDO del PDF.
    Se usa como respaldo para la cédula (si el nombre del archivo no la trae)
    y SIEMPRE para el nombre de la persona, ya que cada entidad (Procuraduría,
    Policía, Contraloría, RNMC, REDAM) redacta la identificación de forma
    distinta: unas veces el nombre va antes de la cédula, otras después,
    y algunas (Contraloría, REDAM) no incluyen nombre en absoluto.
    Retorna (cedula: str|None, nombre: str|None).
    """
    try:
        archivo_pdf.seek(0)
        with pdfplumber.open(archivo_pdf) as pdf:
            texto_completo = ""
            for p in pdf.pages[:3]:
                texto_completo += (p.extract_text() or "") + " "

        tn = normalizar(texto_completo)

        # ── Cédula (respaldo, por si el nombre del archivo no la trae) ────
        cedula = None
        patrones_cedula = [
            r'CEDULA\s+DE\s+CIUDADANIA.{0,40}?(\d{6,12})',
            r'NUMERO\s+DE\s+IDENTIFICACION\s+CC\s*(\d{6,12})',
            r'IDENTIFICACION\s+CC\s*(\d{6,12})',
            r'\bCC\s+(\d{6,12})\b',
        ]
        for patron in patrones_cedula:
            m = re.search(patron, tn)
            if m:
                c = re.sub(r'[^\d]', '', m.group(1))
                if 6 <= len(c) <= 12:
                    cedula = c
                    break

        # ── Nombre (independiente de dónde/cómo aparezca la cédula) ───────
        nombre = None
        patrones_nombre = [
            r'SENOR\(A\)\s+([A-Z ]{8,80}?)\s+IDENTIFICAD',
            r'APELLIDOS\s+Y\s+NOMBRES?[:\s]+([A-Z ]{8,80}?)(?:\s+NO\s|\s*$)',
            r'Y\s+NOMBRE[:\s]+([A-Z ]{8,80}?)\.',
            r'([A-Z ]{8,80}?)\s*,?\s*IDENTIFICAD[OA]\s+CON',
        ]
        for patron in patrones_nombre:
            m = re.search(patron, tn)
            if m:
                cand = re.sub(r'\s+', ' ', m.group(1)).strip()
                cand = re.sub(r'^(EL|LA|SENOR|SENORA|SR|SRA)\s+', '', cand).strip()
                if len(cand) >= 8:
                    nombre = cand
                    break

        return cedula, nombre
    except Exception:
        return None, None


def agrupar_pdfs_por_cedula(archivos_pdf):
    """
    Agrupa los PDFs por cédula. Prioridad:
      1. Cédula tomada del NOMBRE DEL ARCHIVO (rápido y confiable,
         ya que así es como se descargan y guardan estos certificados).
      2. Si el nombre del archivo no trae una cédula válida, cae de
         respaldo al análisis del contenido del PDF.
    El NOMBRE de la persona siempre se busca en el contenido: basta con
    que UNO de los 6 PDFs de esa cédula lo traiga (ej. Procuraduría o
    Policía) para que quede asociado a todo el grupo, aunque otros
    certificados (Contraloría, REDAM) no incluyan nombre.
    Los PDFs sin cédula detectada (ni por archivo ni por contenido)
    quedan en el grupo 'SIN_CEDULA'.
    """
    grupos = {}
    sin_cedula = []

    for pdf in archivos_pdf:
        pdf.seek(0)

        # 1. Intentar por nombre del archivo primero
        cedula = extraer_cedula_de_nombre_archivo(pdf.name)

        # 2. El nombre de la persona siempre se busca en el contenido
        cedula_contenido, nombre_persona = extraer_persona_pdf(pdf)
        pdf.seek(0)

        # 3. Respaldo: si el archivo no traía cédula válida, usar la del contenido
        if not cedula:
            cedula = cedula_contenido

        contenido = pdf.read()

        if not cedula:
            sin_cedula.append((pdf.name, contenido))
            continue

        if cedula not in grupos:
            grupos[cedula] = {'nombre': nombre_persona, 'archivos': []}
        if nombre_persona and not grupos[cedula]['nombre']:
            grupos[cedula]['nombre'] = nombre_persona
        grupos[cedula]['archivos'].append((pdf.name, contenido))

    if sin_cedula:
        grupos['SIN_CEDULA'] = {'nombre': None, 'archivos': sin_cedula}

    return grupos


def fusionar_pdfs_bytes(lista_bytes: list) -> bytes:
    """Fusiona una lista de PDFs (bytes) en uno solo, en el orden dado."""
    writer = pypdf.PdfWriter()
    for contenido in lista_bytes:
        reader = pypdf.PdfReader(BytesIO(contenido))
        for pagina in reader.pages:
            writer.add_page(pagina)

    out = BytesIO()
    writer.write(out)
    out.seek(0)
    return out.getvalue()


def generar_zip_fusionados(grupos: dict, progreso_widget) -> bytes:
    """Genera un ZIP con un PDF fusionado por cada cédula."""
    zip_buf = BytesIO()
    total = len(grupos)

    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for idx, (cedula, info) in enumerate(grupos.items()):
            progreso_widget.progress(
                (idx + 1) / total,
                text=f"Fusionando {idx+1}/{total}: {cedula}"
            )
            bytes_lista = [b for _, b in info['archivos']]
            pdf_fusionado = fusionar_pdfs_bytes(bytes_lista)

            slug = slug_nombre(info['nombre']) if info['nombre'] else cedula
            nombre_archivo = f"{cedula}_{slug}.pdf" if info['nombre'] else f"{cedula}.pdf"

            zf.writestr(nombre_archivo, pdf_fusionado)

    zip_buf.seek(0)
    return zip_buf.getvalue()


# ══════════════════════════════════════════════════════════════════
#  INTERFAZ — TABS
# ══════════════════════════════════════════════════════════════════

st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

tab1, tab2, tab3 = st.tabs([
    "📋  Copiar Obligaciones al Combinado",
    "✂️  Dividir Combinado por Persona",
    "🔗  Unir PDFs por Cédula"
])

# ─────────────────────────────────────────────────────────────────
# TAB 1 — COPIAR OBLIGACIONES (código original conservado)
# ─────────────────────────────────────────────────────────────────
with tab1:
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    col1, col2 = st.columns(2, gap="large")

    with col1:
        st.markdown('<div class="seccion-card"><div class="seccion-card-title">📄 Subir PDFs</div>', unsafe_allow_html=True)
        archivos_pdf = st.file_uploader("PDFs", type=["pdf"], accept_multiple_files=True,
                                        label_visibility="collapsed", key="uploader_pdf_tab1")
        if archivos_pdf:
            st.markdown(f'<div class="chip-row"><span class="chip chip-green">✔ {len(archivos_pdf)} archivo(s)</span></div>',
                        unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="seccion-card"><div class="seccion-card-title">📁 Subir Word combinado</div>', unsafe_allow_html=True)
        archivo_word = st.file_uploader("Word", type=["docx"], label_visibility="collapsed",
                                        key="uploader_word_tab1")
        if archivo_word:
            st.markdown('<div class="chip-row"><span class="chip chip-blue">✔ Word cargado</span></div>',
                        unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    if archivos_pdf and archivo_word:
        cb1, cb2, _ = st.columns([1, 1, 2])
        with cb1:
            btn_diag = st.button("🔍 Diagnosticar", key="btn_diag")
        with cb2:
            btn_proc = st.button("🚀 Procesar documentos", key="btn_proc")

        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

        # DIAGNÓSTICO
        if btn_diag:
            st.subheader("📄 Nombres extraídos de los PDFs")
            nombres_pdf = []
            for pdf in archivos_pdf:
                nombre = extraer_nombre_pdf(pdf)
                if nombre:
                    st.success(f"**{pdf.name}** → `{nombre}`")
                    nombres_pdf.append(nombre)
                else:
                    st.error(f"**{pdf.name}** → ❌ No se detectó nombre")

            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            st.subheader("📋 Contratistas en el Word")
            doc_diag = Document(archivo_word)
            mapa = construir_mapa_word(doc_diag)
            st.markdown(
                f'<div class="chip-row"><span class="chip chip-blue">Bloques encontrados: {len(mapa)}</span></div>',
                unsafe_allow_html=True)

            multiples = {n: v for n, v in mapa.items() if len(v["tablas"]) > 1}
            if multiples:
                st.warning(f"⚠️ {len(multiples)} persona(s) con obligaciones en varias tablas → se fusionarán.")

            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            st.subheader("🔎 Coincidencias PDF ↔ Word")
            ok = 0

            for nombre_pdf in nombres_pdf:
                nombre_word, score = buscar_mejor_match(nombre_pdf, mapa)
                if nombre_word:
                    n_tablas = len(mapa[nombre_word]["tablas"])
                    label = ("exacto" if score == 100 else
                             "alto" if score >= 90 else
                             "bueno" if score >= 75 else "aceptable")
                    st.success(
                        f"✅ **{nombre_pdf}**  \n"
                        f"↳ Word: `{nombre_word}` — confianza {score}/100 ({label}) — {n_tablas} tabla(s)")
                    ok += 1
                else:
                    st.error(f"❌ **{nombre_pdf}** → Sin coincidencia (mejor score: {score}/100)")
                    candidatos = sorted(
                        [(nw, score_nombre(nombre_pdf, nw)) for nw in mapa],
                        key=lambda x: -x[1])[:3]
                    for cand, csc in candidatos:
                        if csc > 0:
                            css = ("match-high" if csc >= 70 else
                                   "match-med" if csc >= 40 else "match-low")
                            st.markdown(
                                f'<div class="match-box {css}">'
                                f'Candidato: <b>{cand}</b> — score {csc}/100</div>',
                                unsafe_allow_html=True)

            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            total_pdf = len(nombres_pdf)
            color_chip = ("chip-green" if ok == total_pdf else
                          "chip-yellow" if ok > 0 else "chip-red")
            st.markdown(
                f'<div class="chip-row"><span class="chip {color_chip}">Coincidencias: {ok} / {total_pdf}</span></div>',
                unsafe_allow_html=True)

        # PROCESAMIENTO
        if btn_proc:
            reporte = []
            procesados = 0
            errores = 0
            no_encontrados = []

            doc = Document(archivo_word)
            mapa_word = construir_mapa_word(doc)
            progreso = st.progress(0, text="Iniciando…")
            total = len(archivos_pdf)

            for idx, pdf in enumerate(archivos_pdf):
                nombre_archivo = pdf.name
                try:
                    nombre_pdf = extraer_nombre_pdf(pdf)
                    if not nombre_pdf:
                        reporte.append({"PDF": nombre_archivo, "Estado": "❌ ERROR",
                                        "Detalle": "No se detectó nombre en el PDF"})
                        errores += 1
                        progreso.progress((idx+1)/total, text=f"{idx+1}/{total} — {nombre_archivo}")
                        continue

                    tabla_pdf = extraer_tabla_obligaciones_pdf(pdf)
                    if not tabla_pdf:
                        reporte.append({"PDF": nombre_archivo, "Estado": "❌ ERROR",
                                        "Detalle": f"No se encontró tabla de obligaciones | {nombre_pdf}"})
                        errores += 1
                        progreso.progress((idx+1)/total, text=f"{idx+1}/{total} — {nombre_archivo}")
                        continue

                    nombre_word, score = buscar_mejor_match(nombre_pdf, mapa_word)
                    if not nombre_word:
                        reporte.append({"PDF": nombre_archivo, "Estado": "❌ ERROR",
                                        "Detalle": f"Sin coincidencia en Word para '{nombre_pdf}' (mejor score: {score}/100)"})
                        no_encontrados.append(nombre_pdf)
                        errores += 1
                        progreso.progress((idx+1)/total, text=f"{idx+1}/{total} — {nombre_archivo}")
                        continue

                    n_tablas = reemplazar_tablas_obligaciones(doc, mapa_word[nombre_word], tabla_pdf)
                    n_obl = len(tabla_pdf) - 1
                    nota = "" if score == 100 else f" [fuzzy {score}/100 → '{nombre_word}']"
                    procesados += 1
                    reporte.append({
                        "PDF": nombre_archivo,
                        "Estado": "✅ OK",
                        "Detalle": f"{nombre_pdf} | {n_obl} obligaciones copiadas | {n_tablas} tabla(s){nota}"
                    })

                except Exception as e:
                    errores += 1
                    reporte.append({"PDF": nombre_archivo, "Estado": "❌ ERROR",
                                    "Detalle": str(e)})

                progreso.progress((idx+1)/total, text=f"Procesando {idx+1}/{total}…")

            output = BytesIO()
            doc.save(output)
            output.seek(0)

            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            cr1, cr2 = st.columns(2)
            with cr1:
                st.success(f"✅ Procesados correctamente: **{procesados}** / {total}")
            with cr2:
                if errores:
                    st.error(f"❌ Errores: **{errores}**")
                else:
                    st.success("Sin errores.")

            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            st.subheader("📊 Reporte de procesamiento")
            st.dataframe(pd.DataFrame(reporte), use_container_width=True)

            if no_encontrados:
                st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
                st.error("⚠️ Contratistas no encontrados en el Word:")
                for p in no_encontrados:
                    st.markdown(f'<span class="chip chip-red">• {p}</span>',
                                unsafe_allow_html=True)

            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            st.download_button(
                label="⬇️  Descargar Word Final",
                data=output,
                file_name="INFORME_FINAL.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    else:
        st.markdown(
            '<div style="text-align:center;padding:2rem 0;color:#8b949e;font-size:1rem;">'
            '📂 &nbsp; Sube los PDFs y el Word combinado para comenzar.'
            '</div>',
            unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────
# TAB 2 — DIVISOR DE INFORMES (código original conservado)
# ─────────────────────────────────────────────────────────────────
with tab2:
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown(
        '<div style="color:#8b949e;font-size:.92rem;margin-bottom:1rem;">'
        'Sube el Word combinado (con todos los informes), detecta automáticamente cada persona '
        'y descarga los informes individuales en ZIP — como Word o PDF.'
        '</div>', unsafe_allow_html=True)

    st.markdown('<div class="seccion-card"><div class="seccion-card-title">📁 Subir Word combinado</div>',
                unsafe_allow_html=True)
    archivo_combinado = st.file_uploader(
        "Combinado", type=["docx"], label_visibility="collapsed",
        key="uploader_combinado_tab2")
    if archivo_combinado:
        st.markdown('<div class="chip-row"><span class="chip chip-blue">✔ Combinado cargado</span></div>',
                    unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    if archivo_combinado:
        btn_detectar = st.button("🔍 Detectar informes", key="btn_detectar")

        if btn_detectar or st.session_state.get('bloques_detectados'):

            if btn_detectar:
                with st.spinner("Analizando el combinado…"):
                    archivo_combinado.seek(0)
                    doc_split = Document(archivo_combinado)
                    bloques = encontrar_bloques_informes(doc_split)
                    st.session_state['bloques_detectados'] = bloques
                    archivo_combinado.seek(0)
                    st.session_state["doc_split_bytes"] = archivo_combinado.read()

            bloques = st.session_state.get('bloques_detectados', [])

            if not bloques:
                st.error("❌ No se encontraron informes en el documento. Verifica que el combinado tenga el formato correcto.")
            else:
                # ── Vista previa ──────────────────────────────────────────
                sin_nombre = sum(1 for b in bloques if not b.get('nombre'))
                col_a, col_b, col_c = st.columns(3)
                with col_a:
                    st.markdown(
                        f'<div class="chip-row"><span class="chip chip-green">📄 {len(bloques)} informes detectados</span></div>',
                        unsafe_allow_html=True)
                with col_b:
                    if sin_nombre:
                        st.markdown(
                            f'<div class="chip-row"><span class="chip chip-yellow">⚠️ {sin_nombre} sin nombre detectado</span></div>',
                            unsafe_allow_html=True)
                with col_c:
                    st.markdown(
                        f'<div class="chip-row"><span class="chip chip-purple">✂️ Listo para dividir</span></div>',
                        unsafe_allow_html=True)

                st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
                st.subheader("📋 Informes detectados")
                st.markdown(
                    '<div style="color:#8b949e;font-size:.85rem;margin-bottom:.6rem;">'
                    'Desmarca los informes que NO quieres incluir en el ZIP.</div>',
                    unsafe_allow_html=True)

                # Checkboxes de selección
                seleccion = {}
                for idx, bloque in enumerate(bloques):
                    nombre = bloque.get('nombre') or f"⚠️ SIN NOMBRE (informe {idx+1})"
                    nombre_key = bloque.get('nombre') or f"SIN_NOMBRE_{idx+1}"
                    elementos = bloque['fin'] - bloque['inicio']
                    estado_css = "split-row-ok" if bloque.get('nombre') else "split-row-warn"

                    col_chk, col_info = st.columns([0.06, 0.94])
                    with col_chk:
                        checked = st.checkbox("", value=True,
                                              key=f"chk_{idx}",
                                              label_visibility="collapsed")
                    with col_info:
                        st.markdown(
                            f'<div class="split-row {estado_css}">'
                            f'<span class="split-num">{idx+1:02d}</span>'
                            f'<span style="flex:1;font-weight:600;">{nombre}</span>'
                            f'<span style="color:#8b949e;font-size:.8rem;">{elementos} elementos</span>'
                            f'</div>',
                            unsafe_allow_html=True)
                    seleccion[nombre_key] = checked

                nombres_seleccionados = [
                    bloque.get('nombre') or f"SIN_NOMBRE_{idx+1}"
                    for idx, bloque in enumerate(bloques)
                    if seleccion.get(bloque.get('nombre') or f"SIN_NOMBRE_{idx+1}", True)
                ]

                n_sel = len(nombres_seleccionados)
                st.markdown(
                    f'<div class="chip-row"><span class="chip chip-blue">Seleccionados: {n_sel} / {len(bloques)}</span></div>',
                    unsafe_allow_html=True)

                st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

                # ── Botones de descarga ───────────────────────────────────
                st.subheader("⬇️ Descargar informes")
                st.markdown(
                    '<div style="color:#8b949e;font-size:.85rem;margin-bottom:.8rem;">'
                    'Elige el formato. Se genera un ZIP con un archivo por persona, '
                    'numerados y nombrados automáticamente.</div>',
                    unsafe_allow_html=True)

                dl_col1, dl_col2 = st.columns(2, gap="large")

                with dl_col1:
                    btn_zip_word = st.button(
                        "📦 Generar ZIP — Word (.docx)",
                        key="btn_zip_word",
                        use_container_width=True)

                with dl_col2:
                    btn_zip_pdf = st.button(
                        "📦 Generar ZIP — PDF",
                        key="btn_zip_pdf",
                        use_container_width=True)

                # ── Generar ZIP Word ──────────────────────────────────────
                if btn_zip_word:
                    if n_sel == 0:
                        st.warning("No hay informes seleccionados.")
                    else:
                        progreso_zip = st.progress(0, text="Preparando…")
                        doc_bytes = st.session_state.get('doc_split_bytes', b'')
                        bloques_act = st.session_state.get('bloques_detectados', [])

                        zip_bytes = generar_zip_informes(
                            doc_bytes, bloques_act, nombres_seleccionados,
                            'docx', progreso_zip)

                        progreso_zip.progress(1.0, text="✅ ZIP listo")
                        st.success(f"✅ ZIP generado con {n_sel} informes Word.")
                        st.download_button(
                            label=f"⬇️ Descargar ZIP ({n_sel} Word)",
                            data=zip_bytes,
                            file_name="INFORMES_INDIVIDUALES_WORD.zip",
                            mime="application/zip",
                            use_container_width=True,
                            key="dl_zip_word"
                        )

                # ── Generar ZIP PDF ───────────────────────────────────────
                if btn_zip_pdf:
                    if n_sel == 0:
                        st.warning("No hay informes seleccionados.")
                    elif not _buscar_libreoffice():
                        st.error(
                            "❌ **LibreOffice no está instalado** en este equipo.  \n"
                            "Descárgalo gratis en [libreoffice.org](https://www.libreoffice.org/download/download-libreoffice/) "
                            "e instálalo, luego reinicia la app.  \n"
                            "También puedes usar el **ZIP Word** y convertir después.")
                    else:
                        st.info("⏳ Convirtiendo a PDF con LibreOffice… esto puede tomar 1-2 minutos para 50 informes.")
                        progreso_zip = st.progress(0, text="Preparando…")
                        doc_bytes = st.session_state.get('doc_split_bytes', b'')
                        bloques_act = st.session_state.get('bloques_detectados', [])

                        zip_bytes = generar_zip_informes(
                            doc_bytes, bloques_act, nombres_seleccionados,
                            'pdf', progreso_zip)

                        progreso_zip.progress(1.0, text="✅ ZIP listo")
                        st.success(f"✅ ZIP generado con {n_sel} informes PDF.")
                        st.download_button(
                            label=f"⬇️ Descargar ZIP ({n_sel} PDFs)",
                            data=zip_bytes,
                            file_name="INFORMES_INDIVIDUALES_PDF.zip",
                            mime="application/zip",
                            use_container_width=True,
                            key="dl_zip_pdf"
                        )
    else:
        st.markdown(
            '<div style="text-align:center;padding:2rem 0;color:#8b949e;font-size:1rem;">'
            '📂 &nbsp; Sube el Word combinado para detectar y dividir los informes.'
            '</div>',
            unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────
# TAB 3 — UNIR PDFS POR CÉDULA (nuevo)
# ─────────────────────────────────────────────────────────────────
with tab3:
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown(
        '<div style="color:#8b949e;font-size:.92rem;margin-bottom:1rem;">'
        'Sube todos los PDFs de antecedentes disciplinarios (ej. 300 PDFs = 50 personas x 6 certificados). '
        'El sistema identifica a cada persona por la cédula que aparece en el NOMBRE del archivo '
        '(ej. <code>42784175_PROCURADURIA.pdf</code>) — así es como normalmente se descargan y guardan estos '
        'certificados. Si algún archivo no trae la cédula en su nombre, el sistema intenta detectarla '
        'leyendo el contenido como respaldo. El nombre de la persona se obtiene automáticamente del '
        'contenido de cualquiera de sus 6 certificados, y se usa para nombrar el PDF final combinado.'
        '</div>', unsafe_allow_html=True)

    st.markdown('<div class="seccion-card"><div class="seccion-card-title">📄 Subir PDFs</div>',
                unsafe_allow_html=True)
    archivos_pdf_tab3 = st.file_uploader(
        "PDFs por cédula", type=["pdf"], accept_multiple_files=True,
        label_visibility="collapsed", key="uploader_pdf_tab3")
    if archivos_pdf_tab3:
        st.markdown(f'<div class="chip-row"><span class="chip chip-green">✔ {len(archivos_pdf_tab3)} archivo(s)</span></div>',
                    unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    if archivos_pdf_tab3:
        btn_agrupar = st.button("🔍 Detectar y agrupar por cédula", key="btn_agrupar")

        if btn_agrupar or st.session_state.get('grupos_cedula'):
            if btn_agrupar:
                with st.spinner("Leyendo cada PDF y detectando cédulas…"):
                    st.session_state['grupos_cedula'] = agrupar_pdfs_por_cedula(archivos_pdf_tab3)

            grupos = st.session_state.get('grupos_cedula', {})
            grupos_validos = {c: v for c, v in grupos.items() if c != 'SIN_CEDULA'}
            n_sin_cedula = len(grupos.get('SIN_CEDULA', {}).get('archivos', []))

            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

            col_a, col_b, col_c = st.columns(3)
            with col_a:
                st.markdown(
                    f'<div class="chip-row"><span class="chip chip-green">👤 {len(grupos_validos)} personas detectadas</span></div>',
                    unsafe_allow_html=True)
            with col_b:
                sin_nombre = sum(1 for v in grupos_validos.values() if not v['nombre'])
                if sin_nombre:
                    st.markdown(
                        f'<div class="chip-row"><span class="chip chip-yellow">⚠️ {sin_nombre} sin nombre en ningún PDF</span></div>',
                        unsafe_allow_html=True)
            with col_c:
                if n_sin_cedula:
                    st.markdown(
                        f'<div class="chip-row"><span class="chip chip-red">❌ {n_sin_cedula} PDF(s) sin cédula detectada</span></div>',
                        unsafe_allow_html=True)

            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            st.subheader("📋 Personas y sus certificados")

            for cedula, info in grupos_validos.items():
                n_archivos = len(info['archivos'])
                nombre_mostrar = info['nombre'] or "⚠️ Nombre no detectado en ningún PDF del grupo"
                css = "match-high" if info['nombre'] else "match-med"
                st.markdown(
                    f'<div class="match-box {css}">'
                    f'<b>{cedula}</b> — {nombre_mostrar} — {n_archivos} PDF(s)'
                    f'</div>', unsafe_allow_html=True)

            if n_sin_cedula:
                st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
                st.error(
                    f"⚠️ {n_sin_cedula} PDF(s) no tenían la frase 'identificado con cédula de "
                    "ciudadanía' detectable (posiblemente PDFs escaneados sin texto seleccionable). "
                    "Revísalos manualmente:")
                for nombre_archivo, _ in grupos.get('SIN_CEDULA', {}).get('archivos', []):
                    st.markdown(f'<span class="chip chip-red">• {nombre_archivo}</span>',
                                unsafe_allow_html=True)

            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

            if st.button("📦 Generar ZIP de PDFs fusionados", key="btn_zip_cedula", use_container_width=True):
                if not grupos_validos:
                    st.warning("No hay grupos con cédula detectada para fusionar.")
                else:
                    progreso = st.progress(0, text="Preparando…")
                    zip_bytes = generar_zip_fusionados(grupos_validos, progreso)
                    progreso.progress(1.0, text="✅ ZIP listo")
                    st.success(f"✅ ZIP generado con {len(grupos_validos)} PDF(s) fusionados.")
                    st.download_button(
                        label="⬇️ Descargar ZIP de PDFs fusionados",
                        data=zip_bytes,
                        file_name="ANTECEDENTES_FUSIONADOS_POR_CEDULA.zip",
                        mime="application/zip",
                        use_container_width=True,
                        key="dl_zip_cedula"
                    )
    else:
        st.markdown(
            '<div style="text-align:center;padding:2rem 0;color:#8b949e;font-size:1rem;">'
            '📂 &nbsp; Sube los PDFs de antecedentes para agruparlos por cédula.'
            '</div>', unsafe_allow_html=True)
